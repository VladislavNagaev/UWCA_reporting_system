import docx
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd
from pandas import DataFrame
import numpy as np
from numpy import ndarray

from .compile_flight_consolidated_data import get_flight_consolidated_data
from .allocate_ranges import allocate_ranges
from .get_plot import get_plot
from .save_figure_to_docx import save_figure_to_docx
from .save_pd_table_to_docx import save_pd_table_to_docx
from .change_orientation import change_orientation
from .get_frequency import get_frequency



def compile_flight_report(
    document:Document,
    data:DataFrame, 
    frange_names_dict,
    feature_names_dict,
    flight_params,
    params,
    dpi=100,
):
    
    # Загрузка параметров
    feature_list_1 = params.get('feature_list_1')
    feature_list_2 = params.get('feature_list_2')
    tables_params = params.get('tables_params')
    franges_name_list = params.get('franges_name_list')
    remove_forward_flights = params.get('remove_forward_flights')
    error_params = params.get('error_params')
    set_errors = error_params.get('set_errors')

    file_name = 'file_name'
    
    # Выделение диапазонов полета
    franges_data_dict = allocate_ranges(data=data)
    
    if remove_forward_flights == True:
        
        # Удаление данных по участкам горизонтального полета
        forward_flights = franges_data_dict.get('forward_flights')
        for frange_number, frange_data in forward_flights.items():
            data.loc[frange_data.index] = np.nan    

        # Удаление данных по участкам горизонтального полета для всего диапазона эксперимента
        franges_data_dict['full']['full'] = data

    # Номер рисунка
    figure_number = 0
    # Номер таблицы
    table_number = 0
    
    # 0. Титульный лист
    document, figure_number, table_number = add_page_1(document, figure_number, table_number, flight_params)
    
    # 0. Введение
    h1 = document.add_heading('Введение', level=1)
    h1.style = 'Heading 1'

    document, figure_number, table_number = add_page_2(document, figure_number, table_number, flight_params)

    # 1. Исходные данные
    h1 = document.add_heading('Исходные данные', level=1)
    h1.style = 'Heading 1'

    # 1.1 Параметры полета
    h2 = document.add_heading('Параметры полета', level=2)
    h2.style = 'Heading 2'  

    document, figure_number, table_number = add_page_3(document, figure_number, table_number, flight_params)
    
    # 1.2 Распределение исходных параметров
    h2 = document.add_heading('Распределение исходных параметров', level=2)
    h2.style = 'Heading 2' 

    document, figure_number, table_number = add_page_4(
        document=document, 
        figure_number=figure_number, 
        table_number=table_number, 
        flight_params=flight_params, 
        data=data, 
        franges_data_dict=franges_data_dict,
        feature_list_1=feature_list_1,
        frange_names_dict=frange_names_dict,
        feature_names_dict=feature_names_dict,
        dpi=dpi,
    )    
    
    # 2. Обработка исходных данных        
    h1 = document.add_heading('Обработка исходных данных', level=1)
    h1.style = 'Heading 1'           

    # 2.1 Расчет действующих нагрузок (по расчетным формулам)
    h2 = document.add_heading('Расчет действующих нагрузок (по расчетным формулам)', level=2)
    h2.style = 'Heading 2' 

    document, figure_number, table_number = add_page_5(
        document=document, 
        figure_number=figure_number, 
        table_number=table_number, 
        flight_params=flight_params, 
        data=data, 
        franges_data_dict=franges_data_dict,
        feature_list_2=feature_list_2,
        frange_names_dict=frange_names_dict,
        feature_names_dict=feature_names_dict,
        dpi=dpi,
    )    

    # 2.2 Значения максимальных действующих нагрузок 
    document = change_orientation(document=document)

    h2 = document.add_heading('Значения максимальных действующих нагрузок', level=2)
    h2.style = 'Heading 2' 

    consolidated_data = get_flight_consolidated_data(
        tables_params=tables_params,
        franges_name_list=franges_name_list,
        franges_data_dict=franges_data_dict,
        file_name=file_name,
    )

    document, figure_number, table_number = add_page_6(
        document=document, 
        figure_number=figure_number, 
        table_number=table_number, 
        consolidated_data=consolidated_data,
        tables_params=tables_params,
        frange_names_dict=frange_names_dict,
        feature_names_dict=feature_names_dict,
        franges_name_list=franges_name_list,
    )    
    
    document = change_orientation(document=document)

    # Таблица ошибок
    if set_errors == True:

        document, figure_number, table_number = add_page_7(
            document=document, 
            figure_number=figure_number, 
            table_number=table_number, 
            data=data,
            error_params=error_params,
            frange_names_dict=frange_names_dict,
            feature_names_dict=feature_names_dict,
            franges_name_list=franges_name_list,
        )    

    return document


def add_page_1(document, figure_number, table_number, flight_params):
    
    date = flight_params.get('date')
    
    if date is not None:
        # Указание даты полета на титульном листе шаблона
        document.paragraphs[9].text = document.paragraphs[9].text.replace('date', date) 
    
    return document, figure_number, table_number


def add_page_2(document, figure_number, table_number, flight_params):
    
    date = flight_params.get('date')
    notation = flight_params.get('notation')
    conclusions = flight_params.get('conclusions')

    # Добавление текста
    text1 = (
        'В данном отчете представлена обработка данных по полету в рамках ГСИ'
        f' за «{date}» для формирования циклов нагружения и анализа нагруженности'
        ' шасси повышенной проходимости.'
    )
    p = document.add_paragraph(text1)
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

    # Добавление текста
    p = document.add_paragraph('', style='Normal')

    if notation is not None:
        # Добавление текста
        text2 = 'Цель работы: '
        p = document.add_paragraph(text2)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        # Добавление текста
        text3 = notation
        p = document.add_paragraph(text3)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY        
        # Добавление текста
        p = document.add_paragraph('', style='Normal')

    if conclusions is not None:
        # Добавление текста
        text4 = 'Результат испытаний: '
        p = document.add_paragraph(text4)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        # Добавление текста
        text5 = conclusions
        p = document.add_paragraph(text5)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        # Добавление текста
        p = document.add_paragraph('', style='Normal')

    # Добавление разрыва страницы
    document.add_page_break()      
    
    return document, figure_number, table_number


def add_page_3(document, figure_number, table_number, flight_params):
    
    flight_sheet_number = flight_params.get('flight_sheet_number')
    takeoff_weight_of_aircraft = flight_params.get('takeoff_weight_of_aircraft')
    mean_aerodynamic_chord = flight_params.get('mean_aerodynamic_chord')
    air_temperature = flight_params.get('air_temperature')
    wind = flight_params.get('wind')
    back_wind = flight_params.get('back_wind')
    wind_speed = flight_params.get('wind_speed')
    wind_course = flight_params.get('wind_course')
    runway_type = flight_params.get('runway_type')
    ground_strength = flight_params.get('ground_strength')    

    # Параметры полета
    if flight_sheet_number is not None:
        text = f"Номер полетного листа – {str(flight_sheet_number)}"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY 

    if takeoff_weight_of_aircraft is not None:
        text = f"Взлетная масса ВС – {str(takeoff_weight_of_aircraft)} кг"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY   

    if mean_aerodynamic_chord is not None:
        text = f"Центровка – {str(round((mean_aerodynamic_chord*100),2))}% САХ "
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY   

    if air_temperature is not None:
        text = f"Температура наружного воздуха – {str(air_temperature)} С°"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY  

    if wind is not None:
        text = f"Ветер – {str(wind)}"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY  

    if back_wind is not None:
        text = f"Попутный ветер – {str(back_wind)}"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY  

    if wind_speed is not None:
        text = f"Скорость ветра – {str(wind_speed)} м/с"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY 

    if wind_course is not None:
        text = f"Курс ветра – {str(wind_course)} °"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY 

    if runway_type is not None:
        text = f"Типа ВПП – {str(runway_type)}"
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY         

    if (runway_type == 'ГВПП' or runway_type == 'ИВПП и ГВПП') and (ground_strength is not None):
        text = f"Прочность грунта – {str(ground_strength)} кгс/см2"       
        p = document.add_paragraph(text)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY    

    # Добавление разрыва страницы
    document.add_page_break()     
    
    return document, figure_number, table_number


def add_page_4(
    document, 
    figure_number, 
    table_number, 
    flight_params, 
    data, 
    franges_data_dict,
    feature_list_1,
    frange_names_dict,
    feature_names_dict,
    dpi,
):

    # Проход по признакам
    for feature in feature_list_1:       

        # Проверка вхождения признака в массив данных
        if feature in data.columns:        

            # Построение графика 
            figure = get_plot(
                data=data,
                franges_data_dict=franges_data_dict,
                franges_dict='full',
                frange_names_dict=frange_names_dict,
                feature_names_dict=feature_names_dict,
                feature1=feature,
                feature2=None,
                time_feature='Time',
                max_value=True,
                fill_between=True,
                dpi=dpi
            )
            # Сохранение графика в документ
            document = save_figure_to_docx(document=document, figure=figure) 
            # Добавление текста  
            figure_number += 1

            if feature_names_dict.get(feature)[1]:
                p = document.add_paragraph(
                    f'Рисунок {figure_number} – {feature_names_dict.get(feature)[0]}, {feature_names_dict.get(feature)[1]}', 
                    style='Normal')
            else:
                p = document.add_paragraph(
                    f'Рисунок {figure_number} – {feature_names_dict.get(feature)[0]}', 
                    style='Normal')

            p.alignment=WD_ALIGN_PARAGRAPH.CENTER   

            # Добавление текста
            p = document.add_paragraph('', style='Normal')

    # Добавление разрыва страницы
    document.add_page_break()      
    
    return document, figure_number, table_number


def add_page_5(
    document, 
    figure_number, 
    table_number, 
    flight_params,
    data, 
    franges_data_dict,
    feature_list_2,
    frange_names_dict,
    feature_names_dict,
    dpi,
):

    # Проход по признакам
    for feature in feature_list_2:       

        # Проверка вхождения признака в массив данных
        if feature in data.columns:        

            # Построение графика 
            figure = get_plot(
                data=data,
                franges_data_dict=franges_data_dict,
                franges_dict='full',
                frange_names_dict=frange_names_dict,
                feature_names_dict=feature_names_dict,
                feature1=feature,
                feature2=None,
                time_feature='Time',
                max_value=True,
                fill_between=True,
                dpi=dpi
            )
            # Сохранение графика в документ
            document = save_figure_to_docx(document=document, figure=figure) 
            # Добавление текста  
            figure_number += 1

            if feature_names_dict.get(feature)[1]:
                p = document.add_paragraph(
                    f'Рисунок {figure_number} – {feature_names_dict.get(feature)[0]}, {feature_names_dict.get(feature)[1]}', 
                    style='Normal')
            else:
                p = document.add_paragraph(
                    f'Рисунок {figure_number} – {feature_names_dict.get(feature)[0]}', 
                    style='Normal')

            p.alignment=WD_ALIGN_PARAGRAPH.CENTER  

            # Добавление текста
            p = document.add_paragraph('', style='Normal')

    # Добавление разрыва страницы
    document.add_page_break()    
    
    return document, figure_number, table_number


def add_page_6(
    document, 
    figure_number, 
    table_number, 
    consolidated_data,
    tables_params,
    frange_names_dict,
    feature_names_dict,
    franges_name_list,
):

    consolidated_data['feature_name'] = consolidated_data.apply(lambda x: feature_names_dict.get(x['row_name'])[0], axis=1)
    
    
    for table_params in tables_params:
        
        table_name = table_params.get('table_name')
        feature_for_table = table_params.get('feature_for_table')
        tracked_features = table_params.get('tracked_features')
        display_columns = table_params.get('display_columns')
        
        if display_columns is None:
            display_columns = feature_for_table

        display_columns = [
            feature for feature in display_columns if feature in list(set(display_columns)&set(consolidated_data.columns))
        ]
        
        for frange_name in franges_name_list:

            # Список участков диапазона
            frange_data = consolidated_data[consolidated_data['frange_name']==frange_name]

            # Проверка не отсутствия данных в массиве
            if not frange_data.empty:       

                extremum_table = pd.concat(
                    frange_data[(frange_data['extremum']=='max')].apply(
                        lambda x: frange_data[(frange_data['row_name']==x['row_name'])], axis=1
                    ).values
                )
                extremum_table = extremum_table[
                    extremum_table['row_name'].apply(lambda x: x in tracked_features)
                ][display_columns]

                # Подготовка таблицы для записи
                table = extremum_table.applymap(
                    lambda x: str(np.around(x, decimals=1)).replace('.', ',') if type(x) is float else str(x)
                )

                # Список наименований столбцов
                column_names = [
                    f'{feature_names_dict.get(columns)[0]}, {feature_names_dict.get(columns)[1]}' 
                    if feature_names_dict.get(columns)[1] else f'{feature_names_dict.get(columns)[0]}'
                    for columns in table.columns
                ]
                # Добавление текста  
                table_number += 1
                p = document.add_paragraph(
                    f'Таблица {table_number} – {table_name} ({frange_names_dict.get(frange_name)})')
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

                # Сохранение таблицы в документ
                document = save_pd_table_to_docx(
                    document=document, 
                    table=table,
                    columns=column_names,
                    vertical_columns=True,
                )               

                # Добавление разрыва страницы
                document.add_page_break() 
    
    return document, figure_number, table_number


def add_page_7(
    document, 
    figure_number, 
    table_number, 
    data,
    error_params,
    frange_names_dict,
    feature_names_dict,
    franges_name_list,
):

    features = error_params.get('features')
    display_columns = error_params.get('display_columns')


    def __get_param_increment(data:DataFrame, n:int=100) -> DataFrame:
        """
        Скорость роста параметра на последующих n точках
        """
        # Среднее значение на следующих n точках
        rolling = data.sort_index(ascending=False).rolling(window=n).mean()\
            .sort_index(ascending=True)
        # Приращение
        param_increment = (rolling.shift(periods=-1)-data).round(3)

        return param_increment

    def __get_error_array(
        data:DataFrame, 
        feature:str, 
        shift:int, 
        max_increment:float, 
        min_range_time:float, 
    ) -> ndarray:
        
        # Скорость роста параметра на последующих n точках
        a = __get_param_increment(data=data[[feature]], n=shift,)
        # Превышение модуля скорости роста максимального заданного инкремента
        b = a.abs() >= max_increment
        # Массив значений параметра времени для классифицированных событий ошибок
        c = data[b[feature]]['Time'].values
        # Массив значений длительности временных промежутков между классифицированными событиями ошибок
        d = np.roll(c, shift=-1) - c
        # Индексы смены участков с классифицированными событиями ошибок
        e = np.where(np.abs(d) >= min_range_time)[0]+1
        # Список массивов классифицированных событий ошибок разбитых на участки
        f = np.split(ary=c, indices_or_sections=e[:-1])
        # Выбор первого элемента в каждом массиве
        g = np.array([i[0] for i in f if i.shape[0]>0])
        
        error_indexes = data['Time'].apply(lambda x: x in g)
        
        error_array = data[error_indexes]
        error_array['feature'] = feature
        error_array['increment_rate'] = a[error_indexes]
        
        return error_array


    # Максимальное значение инкремента до классификации события ошибкой
    max_increment = 1000
    # Размер участка обнаружения ошибок в секундах
    shift_time = 0.02
    # Минимальный оцениваемого размер участка в секундах
    min_range_time = 0.04

    # Частота дискретизации
    frequency = get_frequency(data=data['Time'].values)
    # Размер участка обнаружения ошибок в точках
    shift = int(shift_time * frequency)


    # Массив событий ошибок
    summary_error_array = pd.DataFrame()

    # Проход по массиву оцениваемых признаков
    for i, feature in enumerate(features[3:5]):
        
        # Получение массива данных, соотвествующих классифицированными событиями ошибок
        error_array = __get_error_array(
            data=data,
            feature=feature,
            shift=shift,
            max_increment=max_increment,
            min_range_time=min_range_time,
        )

        # Обновление массива событий ошибок
        summary_error_array = summary_error_array.append(other=error_array, ignore_index=True)

    # Проверка не отсутствия данных в массиве
    if not summary_error_array.empty:    

        summary_error_array['value'] = summary_error_array[['feature', 'Time']].apply(
            lambda x: summary_error_array[(summary_error_array['Time']==x[1])][x[0]].values[0], axis=1
        )
        summary_error_array['frange'] = summary_error_array['frange'].apply(
            lambda x: frange_names_dict.get(x.split('-')[0])
        )
        summary_error_array['feature_name'] = summary_error_array['feature'].apply(
            lambda x: feature_names_dict.get(x)[0] + ', ' + feature_names_dict.get(x)[-1]
        )

        # Выбор столбцов
        summary_error_array = summary_error_array[display_columns]

        # Подготовка таблицы для записи
        table = summary_error_array.applymap(
            lambda x: str(np.around(x, decimals=1)).replace('.', ',') if type(x) is float else str(x)
        )

        # Список наименований столбцов
        column_names = [
            f'{feature_names_dict.get(columns)[0]}, {feature_names_dict.get(columns)[1]}' 
            if feature_names_dict.get(columns)[1] else f'{feature_names_dict.get(columns)[0]}'
            for columns in table.columns
        ]
        # Добавление текста  
        table_number += 1
        p = document.add_paragraph(
            f'Таблица {table_number} – Зафиксированные ошибки')
        p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

        # Сохранение таблицы в документ
        document = save_pd_table_to_docx(
            document=document, 
            table=table,
            columns=column_names,
            vertical_columns=True,
        )               

    return document, figure_number, table_number

