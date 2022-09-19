import docx
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

import pandas as pd
from pandas import DataFrame
import numpy as np

from .save_pd_table_to_docx import save_pd_table_to_docx
from .change_orientation import change_orientation


# Функция создает сводный отчет
def compile_summary_flight_report(
    document:Document,
    consolidated_data:DataFrame, 
    consolidated_addition_data:DataFrame,
    feature_names_dict:dict,
    frange_names_dict:dict,
    params:dict,
    addition_params:dict,
):
    
    tables_params = params.get('tables_params')
    franges_name_list = params.get('franges_name_list')
    
    tables_addition_params = addition_params.get('tables_params')
    addition_franges_name_list = addition_params.get('franges_name_list')
    materials = addition_params.get('materials')
    
    
    # Номер рисунка
    figure_number = 0
    # Номер таблицы
    table_number = 0

    # 1. Сводный отчет
    h1 = document.add_heading('Сводный отчет', level=1)
    h1.style = 'Heading 1'  
    
    # Изменение ориентации страниц документа с вертикальной на горизонтальную
    document = change_orientation(document=document)

    # 1.1 Экстремумы перегрузок
    h2 = document.add_heading('Экстремумы перегрузок', level=1)
    h2.style = 'Heading 2' 
    
    document, figure_number, table_number = add_page_1(
        document=document, 
        figure_number=figure_number, 
        table_number=table_number, 
        consolidated_data=consolidated_data, 
        tables_params=tables_params, 
        franges_name_list=franges_name_list,
        feature_names_dict=feature_names_dict,
        frange_names_dict=frange_names_dict,
    )
    
    # 1.2 Максимальные напряжения
    h2 = document.add_heading('Максимальные напряжения', level=1)
    h2.style = 'Heading 2'  
    
    document, figure_number, table_number = add_page_2(
        document=document, 
        figure_number=figure_number, 
        table_number=table_number, 
        consolidated_data=consolidated_addition_data, 
        tables_params=tables_addition_params, 
        materials=materials,
        franges_name_list=addition_franges_name_list,
        feature_names_dict=feature_names_dict,
        frange_names_dict=frange_names_dict,
    )
    
    # Изменение ориентации страниц документа с горизонтальной на вертикальную
    document = change_orientation(document=document)
    
    return document



def add_page_1(
    document, 
    figure_number, 
    table_number, 
    consolidated_data, 
    tables_params, 
    franges_name_list, 
    feature_names_dict,
    frange_names_dict,
):
    
    consolidated_data['feature_name'] = consolidated_data.apply(lambda x: feature_names_dict.get(x['row_name'])[0], axis=1)
    
    
    for table_params in tables_params:
        
        table_name = table_params.get('table_name')
        feature_for_table = table_params.get('feature_for_table')
        tracked_features = table_params.get('tracked_features')
        display_columns = table_params.get('display_columns')
        
        if display_columns is None:
            display_columns = feature_for_table

        display_columns = [
            feature for feature in display_columns if feature in list(set(display_columns)&set(consolidated_data.columns))
        ]
            
        for frange_name in franges_name_list:

            # Список участков диапазона
            frange_data = consolidated_data[consolidated_data['frange_name']==frange_name]

            # Проверка не отсутствия данных в массиве
            if not frange_data.empty:    

                extremum_table = pd.concat(
                    frange_data[(frange_data['extremum']=='max')].apply(
                        lambda x: frange_data[(frange_data['row_name']==x['row_name'])], axis=1
                    ).values
                )
                extremum_table = extremum_table[
                    extremum_table['row_name'].apply(lambda x: x in tracked_features)
                ][display_columns]

                file_name_list = frange_data['file_name_list'].values[0]

                # Подготовка таблицы для записи
                table = extremum_table.applymap(
                    lambda x: str(np.around(x, decimals=1)).replace('.', ',') if type(x) is float else str(x)
                )

                # Список наименований столбцов
                column_names = [
                    f'{feature_names_dict.get(columns)[0]}, {feature_names_dict.get(columns)[1]}' 
                    if feature_names_dict.get(columns)[1] else f'{feature_names_dict.get(columns)[0]}'
                    for columns in table.columns
                ]
                # Добавление текста  
                table_number += 1
                p = document.add_paragraph(
                    f'Таблица {table_number} – {table_name} ({frange_names_dict.get(frange_name)})')
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

                # Сохранение таблицы в документ
                document = save_pd_table_to_docx(
                    document=document, 
                    table=table,
                    columns=column_names,
                    vertical_columns=True,
                )     

                # Добавление текста
                text1 = ('При определении экстремумов были использованы следующие полеты:\n' + file_name_list)
                p = document.add_paragraph(text1)
                p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY                    

                # Добавление разрыва страницы
                document.add_page_break() 
    
    return document, figure_number, table_number



def add_page_2(
    document, 
    figure_number, 
    table_number, 
    consolidated_data, 
    tables_params,
    materials,
    franges_name_list, 
    feature_names_dict,
    frange_names_dict,
):

    consolidated_data['value'] = consolidated_data.apply(lambda x: consolidated_data.loc[x.name, x['row_name']], axis=1)
    consolidated_data['material'] = consolidated_data.apply(lambda x: materials.get(x['row_name']), axis=1)
    consolidated_data['feature_name'] = consolidated_data.apply(lambda x: feature_names_dict.get(x['row_name'])[0], axis=1)
    consolidated_data['unit'] = consolidated_data.apply(lambda x: feature_names_dict.get(x['row_name'])[1], axis=1)


    consolidated_data['max_value'] = consolidated_data.apply(lambda x: consolidated_data[
        (consolidated_data['extremum']=='max') & 
        (consolidated_data['row_name']==x['row_name']) & 
        (consolidated_data['frange_name']==x['frange_name'])
    ]['value'].values[0], axis=1)

    consolidated_data['min_value'] = consolidated_data.apply(lambda x: consolidated_data[
        (consolidated_data['extremum']=='min') & 
        (consolidated_data['row_name']==x['row_name']) & 
        (consolidated_data['frange_name']==x['frange_name'])
    ]['value'].values[0], axis=1)


    consolidated_data['max_file_name'] = consolidated_data.apply(lambda x: consolidated_data[
        (consolidated_data['extremum']=='max') & 
        (consolidated_data['row_name']==x['row_name']) & 
        (consolidated_data['frange_name']==x['frange_name'])
    ]['file_name'].values[0], axis=1)

    consolidated_data['min_file_name'] = consolidated_data.apply(lambda x: consolidated_data[
        (consolidated_data['extremum']=='min') & 
        (consolidated_data['row_name']==x['row_name']) & 
        (consolidated_data['frange_name']==x['frange_name'])
    ]['file_name'].values[0], axis=1)


    consolidated_data['max_time'] = consolidated_data.apply(lambda x: consolidated_data[
        (consolidated_data['extremum']=='max') & 
        (consolidated_data['row_name']==x['row_name']) & 
        (consolidated_data['frange_name']==x['frange_name'])
    ]['Time'].values[0], axis=1)

    consolidated_data['min_time'] = consolidated_data.apply(lambda x: consolidated_data[
        (consolidated_data['extremum']=='min') & 
        (consolidated_data['row_name']==x['row_name']) & 
        (consolidated_data['frange_name']==x['frange_name'])
    ]['Time'].values[0], axis=1)
    
    
    
    for table_params in tables_params:
        
        table_name = table_params.get('table_name')
        feature_for_table = table_params.get('feature_for_table')
        tracked_features = table_params.get('tracked_features')
        display_columns = table_params.get('display_columns')
        
        if display_columns is None:
            display_columns = feature_for_table

        display_columns = [
            feature for feature in display_columns if feature in list(set(display_columns)&set(consolidated_data.columns))
        ]
            

        for frange_name in franges_name_list:

            # Список участков диапазона
            frange_data = consolidated_data[consolidated_data['frange_name']==frange_name]

            # Проверка не отсутствия данных в массиве
            if not frange_data.empty:       

                extremum_table = frange_data[(frange_data['extremum']=='max')]
                extremum_table = extremum_table[
                    extremum_table['row_name'].apply(lambda x: x in tracked_features)
                ][display_columns]

                file_name_list = frange_data['file_name_list'].values[0]

                # Подготовка таблицы для записи
                table = extremum_table.applymap(
                    lambda x: str(np.around(x, decimals=1)).replace('.', ',') if type(x) is float else str(x)
                )

                # Список наименований столбцов
                column_names = [
                    f'{feature_names_dict.get(columns)[0]}, {feature_names_dict.get(columns)[1]}' 
                    if feature_names_dict.get(columns)[1] else f'{feature_names_dict.get(columns)[0]}'
                    for columns in table.columns
                ]
                # Добавление текста  
                table_number += 1
                p = document.add_paragraph(
                    f'Таблица {table_number} – {table_name} ({frange_names_dict.get(frange_name)})')
                p.alignment=WD_ALIGN_PARAGRAPH.RIGHT

                # Сохранение таблицы в документ
                document = save_pd_table_to_docx(
                    document=document, 
                    table=table,
                    columns=column_names,
                    vertical_columns=True,
                )     

                # Добавление текста
                text1 = ('При определении экстремумов были использованы следующие полеты:\n' + file_name_list)
                p = document.add_paragraph(text1)
                p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY                    

                # Добавление разрыва страницы
                document.add_page_break() 
    
    return document, figure_number, table_number